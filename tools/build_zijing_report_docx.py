from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "大二下" / "选修课" / "财务舞弊" / "紫晶存储案例分析报告.md"
OUT = ROOT / "大二下" / "选修课" / "财务舞弊" / "紫晶存储案例分析报告.docx"


def set_run_font(run, size=12, bold=False, east_asia="宋体", ascii_font="Times New Roman"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)


def set_paragraph_format(paragraph, first_line=False):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_line:
        fmt.first_line_indent = Pt(24)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_run_font(run, size=12)


def add_runs_with_bold(paragraph, text, size=12, east_asia="宋体", bold_all=False):
    pos = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=size, east_asia=east_asia, bold=bold_all)
        run = paragraph.add_run(match.group(1))
        set_run_font(run, size=size, east_asia=east_asia, bold=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, east_asia=east_asia, bold=bold_all)


def main():
    text = SRC.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    section.footer_distance = Cm(1.5)
    add_page_number(section.footer.paragraphs[0])

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.size = Pt(12)
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    in_refs = False
    for raw in lines:
        line = raw.strip()
        if not line:
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(line[2:])
            set_run_font(run, size=15, bold=True, east_asia="黑体")
            continue

        if line.startswith("## "):
            title = line[3:]
            if title == "参考文献":
                in_refs = True
            p = doc.add_paragraph()
            set_paragraph_format(p)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(title)
            set_run_font(run, size=12, bold=True, east_asia="宋体")
            continue

        if line.startswith("### "):
            p = doc.add_paragraph()
            set_paragraph_format(p)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line[4:])
            set_run_font(run, size=12, bold=True, east_asia="宋体")
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style=None)
            set_paragraph_format(p)
            p.paragraph_format.left_indent = Pt(24)
            p.paragraph_format.first_line_indent = Pt(-12)
            run = p.add_run("• ")
            set_run_font(run, size=12)
            add_runs_with_bold(p, line[2:], size=12)
            continue

        p = doc.add_paragraph()
        set_paragraph_format(p, first_line=not in_refs and not line.startswith("["))
        if line.startswith("["):
            p.paragraph_format.first_line_indent = Pt(0)
        add_runs_with_bold(p, line, size=12)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
